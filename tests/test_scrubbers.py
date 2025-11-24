"""Comprehensive test suite for Bleach metadata scrubber."""

import os
import tempfile
from pathlib import Path
from unittest.mock import MagicMock, patch

import pytest
from PIL import Image

from scrubbers import (
    delete_docx,
    get_file_type,
    read_image,
    scrub_image,
    scrub_text_comments,
)


class TestImageScrubbing:
    """Tests for image metadata scrubbing."""

    def test_scrub_image_creates_file(self):
        """Test that scrub_image creates output file successfully."""
        with tempfile.TemporaryDirectory() as td:
            p = Path(td) / "test.png"
            img = Image.new("RGB", (10, 10), color=(123, 222, 111))
            img.save(p)

            out = Path(td) / "test.scrubbed.png"
            result = scrub_image(p, out)
            assert out.exists()
            assert "Scrubbed" in result
            assert out.stat().st_size > 0

    def test_scrub_image_removes_metadata(self):
        """Test that EXIF data is removed from images."""
        with tempfile.TemporaryDirectory() as td:
            p = Path(td) / "test.jpg"
            # Create image with EXIF data
            img = Image.new("RGB", (100, 100), color=(255, 0, 0))
            exif = img.getexif()
            exif[0x0131] = "Test Camera"  # Make tag
            img.save(p, exif=exif)

            out = Path(td) / "test.scrubbed.jpg"
            scrub_image(p, out)

            # Verify metadata is removed
            clean_img = Image.open(out)
            clean_exif = clean_img.getexif()
            assert len(clean_exif) == 0 or 0x0131 not in clean_exif

    def test_scrub_image_with_backup(self):
        """Test that backup is created when requested."""
        with tempfile.TemporaryDirectory() as td:
            p = Path(td) / "test.png"
            img = Image.new("RGB", (10, 10), color=(100, 100, 100))
            img.save(p)

            out = Path(td) / "test.scrubbed.png"
            scrub_image(p, out, backup=True)

            backup_path = p.with_suffix(p.suffix + ".backup")
            assert backup_path.exists()

    def test_scrub_image_nonexistent_file(self):
        """Test that scrubbing nonexistent file raises error."""
        with pytest.raises(FileNotFoundError):
            scrub_image("/nonexistent/file.jpg")

    def test_read_image_returns_metadata(self):
        """Test that read_image returns expected metadata structure."""
        with tempfile.TemporaryDirectory() as td:
            p = Path(td) / "test.png"
            img = Image.new("RGB", (50, 50), color=(128, 128, 128))
            img.save(p)

            metadata = read_image(p)
            assert "format" in metadata
            assert "mode" in metadata
            assert "size" in metadata
            assert metadata["size"] == (50, 50)
            assert metadata["mode"] == "RGB"


class TestTextScrubbing:
    """Tests for text comment removal."""

    def test_scrub_text_comments_removes_prefixes(self):
        """Test that comment lines are removed correctly."""
        with tempfile.TemporaryDirectory() as td:
            p = Path(td) / "sample.txt"
            p.write_text("# comment line\nactual\n// another comment\nkeep\n")
            out = Path(td) / "sample.scrubbed.txt"
            res = scrub_text_comments(p, out, prefixes=["#", "//"])
            assert out.exists()
            content = out.read_text()
            assert "comment" not in content
            assert "actual" in content
            assert "keep" in content

    def test_scrub_text_default_prefixes(self):
        """Test that default comment prefixes work."""
        with tempfile.TemporaryDirectory() as td:
            p = Path(td) / "sample.txt"
            p.write_text("# hash comment\n// slash comment\n; semicolon comment\nkeep this\n")
            out = Path(td) / "sample.scrubbed.txt"
            scrub_text_comments(p, out)
            content = out.read_text()
            assert "keep this" in content
            assert "hash comment" not in content
            assert "slash comment" not in content
            assert "semicolon comment" not in content

    def test_scrub_text_counts_removed_lines(self):
        """Test that result message includes line count."""
        with tempfile.TemporaryDirectory() as td:
            p = Path(td) / "sample.txt"
            p.write_text("# comment1\nkeep1\n# comment2\nkeep2\n")
            out = Path(td) / "sample.scrubbed.txt"
            result = scrub_text_comments(p, out, prefixes=["#"])
            assert "2 lines removed" in result


class TestDocxScrubbing:
    """Tests for DOCX metadata scrubbing."""

    def test_delete_docx_clears_core_properties(self, tmp_path):
        """Test that DOCX metadata is cleared."""
        try:
            from docx import Document
        except ImportError:
            pytest.skip("python-docx not installed")

        p = tmp_path / "test.docx"
        doc = Document()
        doc.add_paragraph("hello world")
        doc.core_properties.author = "Jane Doe"
        doc.core_properties.title = "Secret Title"
        doc.core_properties.comments = "Sensitive comments"
        doc.save(str(p))

        out = tmp_path / "test.scrubbed.docx"
        result = delete_docx(p, out)
        assert out.exists()
        assert "Scrubbed" in result

        new = Document(str(out))
        assert (new.core_properties.author or "") == ""
        assert (new.core_properties.title or "") == ""
        assert (new.core_properties.comments or "") == ""

    def test_delete_docx_preserves_content(self, tmp_path):
        """Test that document content is preserved after scrubbing."""
        try:
            from docx import Document
        except ImportError:
            pytest.skip("python-docx not installed")

        p = tmp_path / "test.docx"
        doc = Document()
        doc.add_paragraph("Important content")
        doc.add_heading("Test Heading", level=1)
        doc.save(str(p))

        out = tmp_path / "test.scrubbed.docx"
        delete_docx(p, out)

        new = Document(str(out))
        text = "\n".join([para.text for para in new.paragraphs])
        assert "Important content" in text
        assert "Test Heading" in text


class TestFileTypeDetection:
    """Tests for file type detection."""

    def test_get_file_type_image(self):
        """Test file type detection for images."""
        with tempfile.TemporaryDirectory() as td:
            p = Path(td) / "test.png"
            img = Image.new("RGB", (10, 10))
            img.save(p)

            file_type = get_file_type(p)
            assert "mime" in file_type
            assert "image" in file_type["mime"].lower()

    def test_get_file_type_nonexistent(self):
        """Test that nonexistent file raises error."""
        with pytest.raises(FileNotFoundError):
            get_file_type("/nonexistent/file.txt")

    def test_get_file_type_caching(self):
        """Test that file type results are cached."""
        with tempfile.TemporaryDirectory() as td:
            p = Path(td) / "test.txt"
            p.write_text("test content")

            # First call
            result1 = get_file_type(str(p))
            # Second call should use cache
            result2 = get_file_type(str(p))

            assert result1 == result2


class TestErrorHandling:
    """Tests for error handling and edge cases."""

    def test_scrub_directory_raises_error(self):
        """Test that attempting to scrub a directory raises error."""
        with tempfile.TemporaryDirectory() as td:
            with pytest.raises(ValueError, match="directory"):
                scrub_image(td)

    def test_invalid_image_raises_error(self):
        """Test that invalid image file raises error."""
        with tempfile.TemporaryDirectory() as td:
            p = Path(td) / "not_an_image.txt"
            p.write_text("This is not an image")

            with pytest.raises(RuntimeError, match="Image scrubbing failed"):
                scrub_image(p)
